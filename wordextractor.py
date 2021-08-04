from docx import Document
from os import walk
from docx2pdf import convert

class Factuur: # Creates an invoice class with everything needed for Excel admin file.
    def __init__(self):
        self.bedrijf = ""
        self.fnummer = ""
        self.datum = ""
        self.incl = ""
        self.excl = ""
        self.btw = ""
    def __str__(self):
        return f"Bedrijf: {self.bedrijf}\nFactuurnummer: {self.fnummer}\nFactuurdatum: {self.datum}\nTotaalbedrag: {self.incl}\n"
    
def extract(fact, table, vb): # Extracts data needed for Excel admin file and stores it into a class
    for row in table[0].rows:
        fact.bedrijf = row.cells[0].paragraphs[2].text
    fact.fnummer = vb.tables[1].cell(1, 0).text
    fact.datum = vb.tables[1].cell(1, 1).text
    # fact.btw = vb.tables[5].cell(0, 1).text
    # fact.excl = vb.tables[3].cell(0, 1).text
    bedrag = vb.tables[5].cell(1, 1).text
    bedrag = bedrag.replace("-", "00")
    fact.incl = bedrag[2:]


def wordext(): # Sends factuur_list to main
    docs = []
    waar = input("Vul de path van de folder in: ")
    filenames = next(walk(f"{waar}"), (None, None, []))[2]
    for i in filenames:
        if ".docx" in i:
            docs.append(i)
    aantal = len(docs)
    # Making a list with all the invoices as Class objects.
    factuur_list = []
    for file in docs:
        factuur = Factuur()
        path = str(file)
        vb = Document(path)
        table = vb.tables
        extract(factuur, table, vb)
        factuur_list.append(factuur)
    convert(f"{waar}") # Converts all .docx to .pdf
    return factuur_list


